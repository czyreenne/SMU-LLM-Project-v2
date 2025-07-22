import os
import re
import fitz  # PyMuPDF
import docx
from lxml import etree
import zipfile
from io import BytesIO

class DocumentExtractor:
    """
    A class to extract text and metadata from PDF and DOCX files
    """
    
    def __init__(self, file_path):
        self.file_path = file_path
        self._validate_file()
    
    def _validate_file(self):
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        file_ext = os.path.splitext(self.file_path)[1].lower()
        if file_ext not in ['.pdf', '.docx']:
            raise ValueError(f"Unsupported file extension: {file_ext}. Only .pdf and .docx are supported.")
    
    def extract(self):
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_pdf()
        elif file_ext == '.docx':
            return self.extract_docx()
    
    def extract_pdf(self):
        try:
            doc = fitz.open(self.file_path)
            
            full_text = ""
            for page in doc:
                page_text = page.get_text()
                page_text = self._clean_text(page_text)
                full_text += page_text + "\n\n"
            
            filename = os.path.basename(self.file_path)
            metadata = {
                "title": doc.metadata.get("title", "") or filename,
                "author": doc.metadata.get("author", "")
            }
            
            doc.close()
            
            return {
                "text": full_text.strip(),
                "metadata": metadata,
                "footnotes": []
            }
            
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_docx(self):
        try:
            doc = docx.Document(self.file_path)
            
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            filename = os.path.basename(self.file_path)
            metadata = {
                "title": doc.core_properties.title or filename,
                "author": doc.core_properties.author or ""
            }
            
            footnotes = self._extract_docx_footnotes()
            
            main_text = "\n\n".join(full_text)
            
            if footnotes:
                footnote_text = "\n\n--FOOTNOTES--\n\n"
                for footnote in footnotes:
                    footnote_text += f"[{footnote['id']}] {footnote['text']}\n"
                
                combined_text = main_text + "\n\n" + footnote_text
            else:
                combined_text = main_text
            
            return {
                "text": combined_text.strip(),
                "metadata": metadata,
                "footnotes": footnotes
            }
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_docx_footnotes(self):
        footnotes = []
        
        try:
            with zipfile.ZipFile(self.file_path) as docx_zip:
                if "word/footnotes.xml" in docx_zip.namelist():
                    with docx_zip.open("word/footnotes.xml") as f:
                        xml_content = f.read()
                        
                        root = etree.fromstring(xml_content)
                        
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        
                        for footnote in root.xpath('//w:footnote', namespaces=ns):
                            footnote_id = footnote.get('{{{w}}}id'.format(**ns), '')
                            if footnote_id not in ['-1', '0']:
                                footnote_text = []
                                for paragraph in footnote.xpath('.//w:p', namespaces=ns):
                                    for text_run in paragraph.xpath('.//w:t', namespaces=ns):
                                        if text_run.text:
                                            footnote_text.append(text_run.text)
                                
                                if footnote_text:
                                    footnotes.append({
                                        'id': footnote_id,
                                        'text': ''.join(footnote_text)
                                    })
        except Exception as e:
            print(f"Warning: Could not extract footnotes: {str(e)}")
            
        return footnotes
    
    def _clean_text(self, text):
        if not text:
            return ""
        
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)
        text = text.replace('fi', 'fi')
        text = text.replace('fl', 'fl')
        
        return text.strip()


# # For testing
# if __name__ == "__main__":
#     pdf_file = r"path\document.pdf"
#     try:
#         extractor = DocumentExtractor(pdf_file)
#         pdf_data = extractor.extract()
#         print(f"PDF Title: {pdf_data['metadata']['title']}")
#         print(f"PDF Author: {pdf_data['metadata']['author']}")
#         print(f"PDF Text (first 200 chars): {pdf_data['text'][:200]}...")
#     except Exception as e:
#         print(f"Error with PDF: {str(e)}")
    
#     docx_file = r"path\document.docx"
#     try:
#         extractor = DocumentExtractor(docx_file)
#         docx_data = extractor.extract()
#         print(f"\nDOCX Title: {docx_data['metadata']['title']}")
#         print(f"DOCX Author: {docx_data['metadata']['author']}")
#         print(f"DOCX Text (first 200 chars): {docx_data['text'][:200]}...")
#         print(f"Number of footnotes: {len(docx_data['footnotes'])}")
#     except Exception as e:
#         print(f"Error with DOCX: {str(e)}")